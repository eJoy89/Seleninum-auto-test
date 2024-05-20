from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def log(message):
    print(f"[LOG] {message}")

def get_kyobobook_prod_name():
    log("Setting up ChromeDriver...")
    service = Service(ChromeDriverManager().install())
    options = webdriver.ChromeOptions()
    driver = webdriver.Chrome(service=service, options=options)
    
    try:
        log("Opening Kyobobook application...")
        driver.get('http://www.kyobobook.co.kr/')
        
        log("Waiting for page to load...")
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'prod_name'))
        )

        log("Finding elements with class name 'prod_name'...")
        prod_name_elements = driver.find_elements(By.CLASS_NAME, 'prod_name')
        
        if prod_name_elements:
            log(f"Found {len(prod_name_elements)} elements. Using the first one.")
            prod_name = prod_name_elements[0].text
            log(f"Product name found: {prod_name}")
            return prod_name
        else:
            log("No elements with class name 'prod_name' found.")
            return None

    except Exception as e:
        log(f"An error occurred: {e}")
        log(f"Error details: {type(e).__name__}, {e}")
        return None
    finally:
        log("Closing browser...")
        driver.quit()

def create_word_document(prod_name):
    log("Creating Word document...")
    doc = Document()

    # Add underlined heading with the fetched product name
    def add_underlined_heading(doc, text, level):
        paragraph = doc.add_heading(text, level)
        if level == 0:  # Underline only for the main heading
            run = paragraph.runs[0]
            run.font.underline = True

    add_underlined_heading(doc, prod_name, level=0)

    # Add other headings
    doc.add_heading('제목 크기, H1', level=1)
    doc.add_heading('제목 크기, H2', level=2)
    doc.add_heading('제목 크기, H3', level=3)
    doc.add_heading('제목 크기, H4', level=4)
    doc.add_heading('제목 크기, H5', level=5)
    doc.add_heading('제목 크기, H6', level=6)

    # Align headings to the center
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document with a specified filename
    doc.save('저장하고 싶은 파일명.docx')
    log("Word document saved as '저장하고 싶은 파일명.docx'")

# Main script
prod_name = get_kyobobook_prod_name()
if prod_name:
    create_word_document(prod_name)
else:
    log("Failed to get product name from Kyobobook.")

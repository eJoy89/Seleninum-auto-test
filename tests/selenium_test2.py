from docx import Document

# 이력서 양식 생성
doc = Document()

# 제목
doc.add_heading('履歴書 (りれきしょ)', level=1)

# 기본정보
doc.add_heading('基本情報', level=2)
table = doc.add_table(rows=7, cols=2)
table.cell(0, 0).text = '氏名'
table.cell(1, 0).text = 'ふりがな'
table.cell(2, 0).text = '生年月日'
table.cell(3, 0).text = '性別'
table.cell(4, 0).text = '現住所'
table.cell(5, 0).text = '電話番号'
table.cell(6, 0).text = 'メールアドレス'

# 학력 및 직력
doc.add_heading('学歴・職歴', level=2)
doc.add_paragraph('学歴:')
doc.add_paragraph('年 月 学校名 卒業')
doc.add_paragraph('年 月 学校名 卒業')

doc.add_paragraph('職歴:')
doc.add_paragraph('年 月 会社名 勤務')
doc.add_paragraph('年 月 会社名 勤務')

# 자격 및 면허
doc.add_heading('資格・免許', level=2)
doc.add_paragraph('年 月 資格名')

# 자기 PR
doc.add_heading('自己PR', level=2)
doc.add_paragraph('ここに自己PRを記入してください。')

# 지원 동기
doc.add_heading('志望動機', level=2)
doc.add_paragraph('ここに志望動機を記入してください。')

# 직무 경력서
doc.add_heading('職務経歴書', level=2)
doc.add_paragraph('ここに職務経歴を記入してください。')

# 파일 저장
file_path = '일본_이력서_양식.docx'
doc.save(file_path)

print(f"File saved at: {file_path}")